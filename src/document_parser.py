"""
SOW and RFP document text extraction.

Handles DOCX, PDF, and plain text formats. Extracts structured
sections, tables, and metadata from procurement documents.
"""
from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)

# Common SOW section headers
SECTION_PATTERNS = [
    r"(?i)^\d+\.?\s*(scope of work|deliverables|timeline|pricing|"
    r"terms and conditions|acceptance criteria|security requirements)",
]


@dataclass
class DocumentSection:
    """A single section extracted from a procurement document."""
    title: str
    content: str
    level: int = 1
    metadata: dict = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """Fully parsed procurement document."""
    filename: str
    sections: list[DocumentSection]
    tables: list[list[list[str]]]
    raw_text: str
    word_count: int


class DocumentParser:
    """Extracts structured content from procurement documents.

    Supports DOCX format with section detection based on heading
    styles and regex pattern matching.
    """

    def parse(self, filepath: str | Path) -> ParsedDocument:
        """Parse a document and return structured sections.

        Parameters
        ----------
        filepath : str or Path
            Path to the document file.

        Returns
        -------
        ParsedDocument
            Parsed document with sections, tables, and metadata.
        """
        filepath = Path(filepath)

        if filepath.suffix == ".docx":
            return self._parse_docx(filepath)
        elif filepath.suffix == ".txt":
            return self._parse_text(filepath)
        else:
            raise ValueError(f"Unsupported format: {filepath.suffix}")

    def _parse_docx(self, filepath: Path) -> ParsedDocument:
        """Parse a DOCX file into structured sections."""
        doc = Document(str(filepath))
        sections: list[DocumentSection] = []
        current_section = DocumentSection(title="Preamble", content="")
        raw_parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            raw_parts.append(text)

            # Detect section boundaries
            if para.style.name.startswith("Heading"):
                level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                if current_section.content:
                    sections.append(current_section)
                current_section = DocumentSection(title=text, content="", level=level)
            else:
                current_section.content += text + "\n"

        if current_section.content:
            sections.append(current_section)

        # Extract tables
        tables = []
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            tables.append(rows)

        raw_text = "\n".join(raw_parts)
        return ParsedDocument(
            filename=filepath.name,
            sections=sections,
            tables=tables,
            raw_text=raw_text,
            word_count=len(raw_text.split()),
        )

    def _parse_text(self, filepath: Path) -> ParsedDocument:
        """Parse a plain text file with regex-based section detection."""
        text = filepath.read_text(encoding="utf-8")
        sections = self._detect_sections(text)
        return ParsedDocument(
            filename=filepath.name,
            sections=sections,
            tables=[],
            raw_text=text,
            word_count=len(text.split()),
        )

    @staticmethod
    def _detect_sections(text: str) -> list[DocumentSection]:
        """Use regex to detect numbered sections in plain text."""
        pattern = re.compile(SECTION_PATTERNS[0], re.MULTILINE)
        matches = list(pattern.finditer(text))

        if not matches:
            return [DocumentSection(title="Full Document", content=text)]

        sections = []
        for i, match in enumerate(matches):
            start = match.start()
            end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
            sections.append(DocumentSection(
                title=match.group().strip(),
                content=text[start:end].strip(),
            ))
        return sections
