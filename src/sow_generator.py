"""
SOW document assembly engine.

Combines template sections, matched clauses, and compliance fixes
into a final Statement of Work document.
"""
from __future__ import annotations

import logging
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches

logger = logging.getLogger(__name__)


@dataclass
class SOWMetadata:
    """Metadata for the generated SOW document."""
    title: str
    contract_number: str
    department: str
    effective_date: date
    expiry_date: date
    security_level: str = "Protected B"


class SOWGenerator:
    """Assembles SOW documents from templates and matched clauses.

    Parameters
    ----------
    template_dir : str
        Path to directory containing DOCX templates.
    """

    def __init__(self, template_dir: str = "templates/") -> None:
        self.template_dir = Path(template_dir)

    def generate(
        self,
        metadata: SOWMetadata,
        sections: dict[str, str],
        output_path: str,
    ) -> Path:
        """Generate a complete SOW document.

        Parameters
        ----------
        metadata : SOWMetadata
            Document metadata (title, contract number, etc.).
        sections : dict[str, str]
            Section title → content mapping.
        output_path : str
            Where to save the generated DOCX.

        Returns
        -------
        Path
            Path to the generated document.
        """
        doc = Document()
        self._add_cover_page(doc, metadata)
        self._add_table_of_contents(doc)

        for title, content in sections.items():
            doc.add_heading(title, level=1)
            for paragraph in content.split("\n"):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())

        self._add_signature_block(doc, metadata)

        output = Path(output_path)
        doc.save(str(output))
        logger.info("SOW generated: %s (%d sections)", output, len(sections))
        return output

    def _add_cover_page(self, doc: Document, meta: SOWMetadata) -> None:
        """Add a formatted cover page."""
        doc.add_paragraph("")  # spacer
        title = doc.add_heading(meta.title, level=0)
        title.alignment = 1  # center

        doc.add_paragraph(f"Contract: {meta.contract_number}")
        doc.add_paragraph(f"Department: {meta.department}")
        doc.add_paragraph(f"Effective: {meta.effective_date}")
        doc.add_paragraph(f"Security: {meta.security_level}")
        doc.add_page_break()

    def _add_table_of_contents(self, doc: Document) -> None:
        """Insert a placeholder for Table of Contents."""
        doc.add_heading("Table of Contents", level=1)
        doc.add_paragraph("[Auto-generated on document finalization]")
        doc.add_page_break()

    def _add_signature_block(self, doc: Document, meta: SOWMetadata) -> None:
        """Add signature lines at the end."""
        doc.add_page_break()
        doc.add_heading("Signatures", level=1)
        for party in ["Contracting Authority", "Technical Authority", "Contractor"]:
            doc.add_paragraph(f"{party}: ______________________________")
            doc.add_paragraph(f"Date: ______________________________")
            doc.add_paragraph("")
